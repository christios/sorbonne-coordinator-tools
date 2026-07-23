"""Fill the approved English syllabus template from a stored syllabus record."""

from __future__ import annotations

from copy import deepcopy
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import _Cell, Table


TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "assets" / "syllabus_template_en.docx"


def build_syllabus_docx(syllabus: dict[str, Any], output_path: Path) -> None:
    """Write an editable DOCX populated from the approved English template."""
    document = Document(TEMPLATE_PATH)
    content = _record(syllabus.get("content"))

    _fill_identification(document.tables[0], syllabus, _record(content.get("identification")))
    _fill_contacts(document.tables[1], document.tables[2], _record(content.get("contacts")))
    _set_cell_text(document.tables[3].cell(0, 0), _text(_record(content.get("description")).get("overview")))
    _fill_delivery(document.tables[4], _record(content.get("delivery")))
    _fill_learning_outcomes(document.tables[5], document.tables[6], _record(content.get("learningOutcomes")))
    _fill_schedule(document.tables[7], _rows(content.get("schedule")))
    _fill_bibliography(document.tables[8], _record(content.get("bibliography")))
    _fill_teaching_approach(document, _record(content.get("teachingApproach")))
    _fill_assessment(document, _record(content.get("assessment")), _record(content.get("learningOutcomes")))
    _fill_document_control(document.tables[-1], _record(content.get("documentControl")))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _fill_identification(table: Table, syllabus: dict[str, Any], identification: dict[str, Any]) -> None:
    values = [
        _text(syllabus.get("academicYear")),
        _text(syllabus.get("courseTitle")),
        _text(syllabus.get("courseCode")),
        _text(identification.get("degreeLevelAndSemester")),
        _text(identification.get("programmeTitle")),
        _text(identification.get("ects")),
    ]
    for index, value in enumerate(values):
        _set_cell_text(table.cell(index, 1), value)

    contact_hours = _record(identification.get("contactHours"))
    for row_index, label in enumerate(
        ["Lectures", "Tutorials", "Workshops", "Seminars", "Laboratory", "Other"], start=6
    ):
        _set_cell_text(table.cell(row_index, 3), _text(contact_hours.get(label)))

    _set_cell_text(table.cell(12, 1), _list_or_legacy(identification, "prerequisiteItems", "prerequisites"))
    _set_cell_text(table.cell(13, 1), _list_or_legacy(identification, "equipmentItems", "equipment"))


def _fill_contacts(instructor_table: Table, administrative_table: Table, contacts: dict[str, Any]) -> None:
    instructor = _record(contacts.get("instructor"))
    _set_cell_text(instructor_table.cell(0, 1), _text(instructor.get("Name")))
    _set_cell_text(instructor_table.cell(1, 1), _text(instructor.get("Academic rank / status")))
    _set_cell_text(
        instructor_table.cell(2, 1), _list_or_legacy(instructor, "affiliations", "Affiliation(s)", field="name")
    )
    _set_cell_text(instructor_table.cell(3, 1), _office_hours(instructor))
    _set_cell_text(instructor_table.cell(4, 1), _text(instructor.get("Email")))

    administrative = contacts.get("administrativeContact")
    if isinstance(administrative, str):
        _set_cell_text(administrative_table.cell(0, 1), administrative)
        return
    admin = _record(administrative)
    parts = [
        f"Name: {_text(admin.get('name'))}" if _text(admin.get("name")) else "",
        f"Contact details: {_text(admin.get('contactDetails'))}" if _text(admin.get("contactDetails")) else "",
    ]
    _set_cell_text(administrative_table.cell(0, 1), "\n".join(part for part in parts if part))


def _fill_delivery(table: Table, delivery: dict[str, Any]) -> None:
    mode = _text(delivery.get("mode"))
    _set_cell_text(table.cell(2, 0), "☒" if mode == "Face-to-Face Delivery" else "")
    _set_cell_text(
        table.cell(2, 1), _percentage(delivery.get("faceToFacePercent")) if mode == "Blended Learning Delivery" else ""
    )
    _set_cell_text(
        table.cell(2, 2), _percentage(delivery.get("onlinePercent")) if mode == "Blended Learning Delivery" else ""
    )


def _fill_learning_outcomes(plo_table: Table, clo_table: Table, outcomes: dict[str, Any]) -> None:
    plos = _rows(outcomes.get("plos"))
    _ensure_data_rows(plo_table, header_rows=1, required_rows=max(6, len(plos)))
    for index, row in enumerate(plo_table.rows[1:]):
        source = plos[index] if index < len(plos) else {}
        _set_cell_text(row.cells[0], _text(source.get("code")) or f"PLO {index + 1}")
        outcome = _text(source.get("outcome")) or _strip_prefix(_text(source.get("legacyText")), "PLO")
        _set_cell_text(row.cells[1], outcome)

    clos = _rows(outcomes.get("clos"))
    _ensure_data_rows(clo_table, header_rows=1, required_rows=max(6, len(clos)))
    for index, row in enumerate(clo_table.rows[1:]):
        source = clos[index] if index < len(clos) else {}
        _set_cell_text(row.cells[0], _text(source.get("clo")))
        _set_cell_text(row.cells[1], _text(source.get("plo")))
        _set_cell_text(row.cells[2], _text(source.get("skills")))


def _fill_schedule(table: Table, schedule: list[dict[str, Any]]) -> None:
    _ensure_data_rows(table, header_rows=1, required_rows=max(16, len(schedule)))
    for index, row in enumerate(table.rows[1:]):
        source = schedule[index] if index < len(schedule) else {}
        _set_cell_text(row.cells[0], str(index + 1) if source else "")
        _set_cell_text(row.cells[1], _display_date(source.get("date")))
        _set_cell_text(row.cells[2], _text(source.get("topic")))
        _set_cell_text(row.cells[3], _schedule_learning_details(source))


def _fill_bibliography(table: Table, bibliography: dict[str, Any]) -> None:
    resources = [
        _format_resources(_rows(bibliography.get("books")), "book"),
        _format_resources(_rows(bibliography.get("websites")), "website"),
        _format_resources(_rows(bibliography.get("journalArticles")), "article"),
    ]
    for row, value in zip(table.rows, resources, strict=True):
        _set_cell_text(row.cells[1], value)


def _fill_teaching_approach(document: Document, approach: dict[str, Any]) -> None:
    _set_paragraph_after(document, "8.1. Teaching Methods and Learning Activities", _text(approach.get("methods")))
    _set_paragraph_after(document, "8.2. Student Engagement", _text(approach.get("engagement")))
    _set_paragraph_after(document, "8.3. Feedback and Academic Progress", _text(approach.get("feedback")))


def _fill_assessment(document: Document, assessment: dict[str, Any], learning_outcomes: dict[str, Any]) -> None:
    assessment_table = document.tables[9]
    items = _rows(assessment.get("items"))
    _ensure_data_rows(assessment_table, header_rows=1, required_rows=max(14, len(items)))
    clo_by_id = {row.get("id"): _text(row.get("clo")) for row in _rows(learning_outcomes.get("clos"))}
    for index, row in enumerate(assessment_table.rows[1:]):
        source = items[index] if index < len(items) else {}
        aligned_clos = _text(source.get("clos")) or ", ".join(
            value
            for value in (
                clo_by_id.get(item) for item in source.get("cloIds", []) if isinstance(source.get("cloIds"), list)
            )
            if value
        )
        values = [
            _display_date(source.get("date")),
            _text(source.get("type")),
            _percentage(source.get("weight")),
            aligned_clos,
            _text(source.get("aiPolicy")) or _text(source.get("ai")),
        ]
        for cell, value in zip(row.cells, values, strict=True):
            _set_cell_text(cell, value)

    _fill_ai_policy(document.tables[10], assessment)
    _set_paragraph_after(
        document,
        "Additional Instructions regarding AI where its use is permitted or required",
        _text(assessment.get("aiInstructions")),
    )
    _set_paragraph_after(document, "9.4. Assessment Methodologies:", _text(assessment.get("methodologies")))
    _fill_rubrics(document, _rows(assessment.get("rubrics")))
    _set_paragraph_after(document, "9.6. Late submission policy:", _text(assessment.get("lateSubmissionPolicy")))


def _fill_ai_policy(table: Table, assessment: dict[str, Any]) -> None:
    policy = _text(assessment.get("aiPolicy")) or "AI Prohibited"
    selected_index = {
        "AI Prohibited": 1,
        "AI Permitted as a Support Tool": 2,
        "Restricted AI Use": 3,
        "AI Required": 4,
        "Other (Specify)": 5,
    }.get(policy)
    if selected_index is not None:
        _set_cell_text(table.cell(selected_index, 0), f"☒ {policy}")

    if policy == "AI Prohibited":
        cell = table.cell(1, 1)
        additional = _text(assessment.get("aiVerificationMechanism"))
        _set_cell_text(
            cell,
            _mark_first_checkbox(cell.text)
            + (f"\n\nVerification mechanism used by the instructor: {additional}" if additional else ""),
        )
    elif policy == "AI Permitted as a Support Tool":
        cell = table.cell(2, 1)
        allowed = {item for item in assessment.get("aiAllowedUses", []) if isinstance(item, str)}
        description = cell.text
        for use in allowed:
            description = description.replace(f"☐ {use}", f"☒ {use}")
        other_uses = _list_or_legacy(assessment, "aiOtherUses", "aiOtherUse")
        if other_uses:
            description = description.replace("☐ Other (please specify):", f"☒ Other (please specify): {other_uses}")
        _set_cell_text(cell, description)
    elif policy == "Other (Specify)":
        _set_cell_text(table.cell(5, 1), _text(assessment.get("aiCustomPolicy")))


def _fill_rubrics(document: Document, rubrics: list[dict[str, Any]]) -> None:
    tables = list(document.tables[11:13])
    while len(tables) < len(rubrics):
        clone = deepcopy(tables[-1]._tbl)
        tables[-1]._tbl.addnext(clone)
        tables = list(document.tables[11 : 11 + len(tables) + 1])

    for index, table in enumerate(tables):
        rubric = rubrics[index] if index < len(rubrics) else {}
        criteria = _rows(rubric.get("criteria"))
        _set_cell_text(
            table.cell(0, 0), f"Assessment type: {_text(rubric.get('assignment'))}" if rubric else "Assessment type:"
        )
        _ensure_data_rows(table, header_rows=2, required_rows=max(4, len(criteria)))
        for row_index, row in enumerate(table.rows[2:]):
            criterion = criteria[row_index] if row_index < len(criteria) else {}
            values = [
                criterion.get("criterion"),
                criterion.get("inadequate"),
                criterion.get("meets"),
                criterion.get("exceeds"),
            ]
            for cell, value in zip(row.cells, values, strict=True):
                _set_cell_text(cell, _text(value))


def _fill_document_control(table: Table, control: dict[str, Any]) -> None:
    values = [
        _display_date(control.get("creationDate")),
        _text(control.get("departmentName")),
        _display_date(control.get("approvalDate")),
        _text(control.get("versionNumber")),
        _text(control.get("approver")),
    ]
    for index, value in enumerate(values):
        _set_cell_text(table.cell(index, 1), value)


def _ensure_data_rows(table: Table, *, header_rows: int, required_rows: int) -> None:
    current = len(table.rows) - header_rows
    if current >= required_rows:
        return
    template_row = table.rows[-1]._tr
    for _ in range(required_rows - current):
        table._tbl.append(deepcopy(template_row))


def _set_paragraph_after(document: Document, heading: str, text: str) -> None:
    normalized_heading = heading.rstrip().lower()
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.rstrip().lower() != normalized_heading:
            continue
        for candidate in document.paragraphs[index + 1 :]:
            if candidate.text.strip():
                break
            _set_paragraph_text(candidate, text)
            return


def _set_cell_text(cell: _Cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "")


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    existing_runs = list(paragraph.runs)
    style = deepcopy(existing_runs[0]._r.rPr) if existing_runs and existing_runs[0]._r.rPr is not None else None
    paragraph.clear()
    if not text:
        return
    run = paragraph.add_run(text)
    if style is not None:
        run._r.insert(0, style)


def _format_resources(resources: list[dict[str, Any]], kind: str) -> str:
    return "\n".join(value for value in (_format_resource(entry, kind) for entry in resources) if value)


def _format_resource(entry: dict[str, Any], kind: str) -> str:
    if legacy := _text(entry.get("legacyText")):
        return legacy
    if kind == "book":
        return _sentence(
            [
                _text(entry.get("authors")),
                _year(entry),
                _text(entry.get("title")),
                _text(entry.get("edition")),
                _text(entry.get("publisher")),
                _prefixed("ISBN/URL", entry.get("isbn")),
            ]
        )
    if kind == "website":
        return _sentence(
            [
                _text(entry.get("organisation")),
                _text(entry.get("url")),
                _prefixed("Accessed", _display_date(entry.get("accessedDate"))),
            ]
        )
    return _sentence(
        [
            _text(entry.get("authors")),
            _year(entry),
            _text(entry.get("title")),
            _text(entry.get("journal")),
            _volume_issue(entry),
            _text(entry.get("pages")),
            _text(entry.get("doi")),
        ]
    )


def _sentence(parts: list[str]) -> str:
    return ". ".join(part.strip().rstrip(".") for part in parts if part.strip()) + (
        "." if any(part.strip() for part in parts) else ""
    )


def _year(entry: dict[str, Any]) -> str:
    return f"({_text(entry.get('year'))})" if _text(entry.get("year")) else ""


def _volume_issue(entry: dict[str, Any]) -> str:
    volume, issue = _text(entry.get("volume")), _text(entry.get("issue"))
    if volume and issue:
        return f"{volume}({issue})"
    return volume or issue


def _prefixed(label: str, value: Any) -> str:
    text = _text(value)
    return f"{label}: {text}" if text else ""


def _office_hours(instructor: dict[str, Any]) -> str:
    office_hours = _rows(instructor.get("officeHours"))
    if not office_hours:
        return _text(instructor.get("Office hours and location"))
    entries = []
    for item in office_hours:
        if legacy := _text(item.get("legacyText")):
            entries.append(legacy)
            continue
        time = "–".join(value for value in [_text(item.get("startTime")), _text(item.get("endTime"))] if value)
        summary = " · ".join(value for value in [_text(item.get("day")), time, _text(item.get("location"))] if value)
        if summary:
            entries.append(summary)
    return "\n".join(entries)


def _list_or_legacy(container: dict[str, Any], list_key: str, legacy_key: str, *, field: str = "text") -> str:
    entries = _rows(container.get(list_key))
    if entries:
        return "\n".join(_text(entry.get(field)) for entry in entries if _text(entry.get(field)))
    return _text(container.get(legacy_key))


def _schedule_learning_details(row: dict[str, Any]) -> str:
    parts = []
    if pre_class := _text(row.get("preClass") or row.get("activities")):
        parts.append(f"Pre-class learning activities:\n{pre_class}")
    if assessments := _text(row.get("assessments")):
        parts.append(f"Assessments:\n{assessments}")
    return "\n\n".join(parts)


def _mark_first_checkbox(value: str) -> str:
    return value.replace("☐", "☒", 1)


def _strip_prefix(value: str, prefix: str) -> str:
    return value.removeprefix(prefix).lstrip(" 0123456789:.")


def _percentage(value: Any) -> str:
    text = _text(value)
    return f"{text}%" if text and not text.endswith("%") else text


def _display_date(value: Any) -> str:
    text = _text(value)
    if not text:
        return ""
    try:
        return date.fromisoformat(text).strftime("%d/%m/%Y")
    except ValueError:
        try:
            return datetime.strptime(text, "%d/%m/%Y").strftime("%d/%m/%Y")
        except ValueError:
            return text


def _record(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _rows(value: Any) -> list[dict[str, Any]]:
    if isinstance(value, str):
        return [{"legacyText": value}] if value.strip() else []
    if not isinstance(value, list):
        return []
    return [
        item if isinstance(item, dict) else {"legacyText": item}
        for item in value
        if isinstance(item, dict) or isinstance(item, str) and item.strip()
    ]


def _text(value: Any) -> str:
    return value.strip() if isinstance(value, str) else str(value).strip() if isinstance(value, (int, float)) else ""
