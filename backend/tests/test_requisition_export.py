from docx import Document
from docx.oxml.ns import qn

from sorbonne.services.requisition_export import build_requisition_docx


def _requisition(**overrides) -> dict:
    content = {
        "department": "Department of Sciences and Engineering",
        "program": "Bachelor in Physics",
        "jobTitle": "Part Time Lecturer",
        "classType": "TD",
        "employeeType": "PT",
        "contractFrom": "2026-08-31",
        "contractTo": "2026-12-31",
        "courses": [
            {"subjectCode": "PHY", "courseNumber": "101", "level": "L1", "title": "Mechanics", "hours": "24"},
            {"subjectCode": "PHY", "courseNumber": "102", "level": "L1", "title": "Waves", "hours": "12 TD"},
        ],
    }
    content.update(overrides.pop("content", {}))
    return {"employeeName": "Dr Amira Example", "academicYear": "2026-2027", "content": content, **overrides}


def _row_text(table, row: int) -> str:
    return " ".join(node.text or "" for node in table.rows[row]._tr.iter(qn("w:t")))


def _control_values(table, row: int) -> list[str]:
    controls = table.rows[row]._tr.iter(qn("w:sdt"))
    return [next(control.find(qn("w:sdtContent")).iter(qn("w:t"))).text or "" for control in controls]


def test_fills_the_approved_template_details(tmp_path) -> None:
    output = tmp_path / "requisition.docx"
    build_requisition_docx(_requisition(), output)

    document = Document(output)
    details = document.tables[0]
    assert details.rows[0].cells[1].text == "Dr Amira Example"
    assert details.rows[6].cells[1].text == "36"
    assert _control_values(details, 1) == ["Department of Sciences and Engineering"]
    assert _control_values(details, 2) == ["Bachelor in Physics"]
    assert _control_values(details, 3) == ["Part Time Lecturer"]
    assert _control_values(details, 5) == ["8/31/2026", "12/31/2026"]
    assert _control_values(details, 7) == ["TD"]


def test_keeps_the_template_letterhead_and_approval_sections(tmp_path) -> None:
    output = tmp_path / "requisition.docx"
    build_requisition_docx(_requisition(), output)

    document = Document(output)
    assert [len(table.rows) for table in document.tables] == [9, 3, 1, 3, 4, 2]
    assert "Signatories" in "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert document.tables[3].rows[0].cells[1].text.strip() == "Head of Department"
    assert document.tables[5].rows[0].cells[3].text.strip() == "Date"
    assert document.sections[0].header.paragraphs


def test_writes_one_course_row_per_course_with_the_level_control(tmp_path) -> None:
    output = tmp_path / "requisition.docx"
    build_requisition_docx(_requisition(), output)

    courses = Document(output).tables[1]
    assert [row.cells[0].text for row in courses.rows] == ["Subject Code", "PHY", "PHY"]
    assert [cell.text for cell in courses.rows[1].cells] == ["PHY", "101", "Mechanics", "24"]
    assert _control_values(courses, 1) == ["L1"]
    assert [cell.text for cell in courses.rows[2].cells] == ["PHY", "102", "Waves", "12 TD"]
    identifiers = [element.get(qn("w:val")) for element in courses.rows[1]._tr.iter(qn("w:id"))]
    assert identifiers != [element.get(qn("w:val")) for element in courses.rows[2]._tr.iter(qn("w:id"))]


def test_ticks_the_employee_type_matching_the_requisition(tmp_path) -> None:
    part_time = tmp_path / "part-time.docx"
    full_time = tmp_path / "full-time.docx"
    build_requisition_docx(_requisition(), part_time)
    build_requisition_docx(_requisition(content={"employeeType": "FT"}), full_time)

    assert _control_values(Document(part_time).tables[0], 8) == ["☐", "☒"]
    assert _control_values(Document(full_time).tables[0], 8) == ["☒", "☐"]


def test_keeps_a_blank_course_row_when_no_courses_exist(tmp_path) -> None:
    output = tmp_path / "empty.docx"
    build_requisition_docx(_requisition(content={"courses": []}), output)

    document = Document(output)
    assert [row.cells[0].text for row in document.tables[1].rows] == ["Subject Code", ""]
    assert [cell.text for cell in document.tables[1].rows[1].cells] == ["", "", "", ""]
    assert document.tables[0].rows[6].cells[1].text == "0"
