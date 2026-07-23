from docx import Document

from sorbonne.services.syllabus_export import build_syllabus_docx


EXPANDED_PLO_TABLE_ROW_COUNT = 8
LEGACY_CONTACT = "Name: Mrs Sample Contact\\nContact details: s.contact@sorbonne.ae"


def test_builds_a_filled_template_with_repeatable_course_content(tmp_path) -> None:
    syllabus = {
        "courseTitle": "Climate Change Law and Policy",
        "courseCode": "PublicAffairs585",
        "academicYear": "2026-2027",
        "content": {
            "identification": {
                "degreeLevelAndSemester": "Bachelor 3, Semester 2",
                "programmeTitle": "Bachelor in Public Affairs",
                "ects": "6",
                "contactHours": {"Lectures": "20", "Workshops": "30"},
                "prerequisiteItems": [{"id": "prereq-1", "text": "Introduction to public policy"}],
                "equipmentItems": [{"id": "equipment-1", "text": "Laptop"}],
            },
            "contacts": {
                "instructor": {
                    "Name": "M. Modele Syllabus",
                    "Academic rank / status": "Associate Professor",
                    "affiliations": [{"id": "affiliation-1", "name": "Sorbonne University Abu Dhabi"}],
                    "officeHours": [
                        {
                            "id": "office-hour-1",
                            "day": "Tuesday",
                            "startTime": "10:00",
                            "endTime": "12:00",
                            "location": "A6-117",
                        }
                    ],
                    "Email": "m.syllabus@sorbonne.ae",
                },
                "administrativeContact": {"name": "Mrs Sample Contact", "contactDetails": "s.contact@sorbonne.ae"},
            },
            "description": {"overview": "An applied course on climate governance."},
            "delivery": {"mode": "Blended Learning Delivery", "faceToFacePercent": "70", "onlinePercent": "30"},
            "learningOutcomes": {
                "plos": [{"id": "plo-1", "code": "PLO 1", "outcome": "Evaluate climate policy."}],
                "clos": [
                    {
                        "id": "clo-1",
                        "clo": "CLO 1: Analyse climate law.",
                        "plo": "PLO 1: Evaluate climate policy.",
                        "skills": "Critical reasoning",
                    }
                ],
            },
            "schedule": [
                {
                    "id": "session-1",
                    "date": "2026-09-01",
                    "topic": "Climate governance",
                    "preClass": "Read chapter 1",
                    "assessments": "Short quiz",
                }
            ],
            "bibliography": {
                "books": [
                    {
                        "id": "book-1",
                        "authors": "A. Author",
                        "title": "Climate Law",
                        "year": "2026",
                        "publisher": "Press",
                    }
                ],
                "websites": [
                    {
                        "id": "website-1",
                        "organisation": "UNFCCC",
                        "url": "https://unfccc.int",
                        "accessedDate": "2026-07-23",
                    }
                ],
                "journalArticles": [
                    {
                        "id": "article-1",
                        "authors": "B. Author",
                        "title": "Climate governance",
                        "journal": "Policy Review",
                        "year": "2026",
                        "doi": "https://doi.org/example",
                    }
                ],
            },
            "teachingApproach": {
                "methods": "Case studies",
                "engagement": "Seminar discussion",
                "feedback": "Written feedback",
            },
            "assessment": {
                "items": [
                    {
                        "id": "assessment-1",
                        "date": "2026-10-10",
                        "type": "Policy brief",
                        "weight": "40",
                        "clos": "CLO 1: Analyse climate law.",
                        "aiPolicy": "AI Prohibited",
                    }
                ],
                "aiPolicy": "AI Prohibited",
                "aiVerificationMechanism": "Oral follow-up discussion",
                "aiInstructions": "Cite all sources.",
                "methodologies": "Written policy analysis",
                "rubrics": [
                    {
                        "id": "rubric-1",
                        "assignment": "Policy brief",
                        "criteria": [
                            {
                                "id": "criterion-1",
                                "criterion": "Analysis",
                                "inadequate": "Limited",
                                "meets": "Sound",
                                "exceeds": "Excellent",
                            }
                        ],
                    }
                ],
                "lateSubmissionPolicy": "Late work is subject to the course policy.",
            },
            "documentControl": {
                "creationDate": "2026-07-23",
                "departmentName": "SCEN",
                "approvalDate": "2026-09-01",
                "versionNumber": "1.0",
                "approver": "Head of Department",
            },
        },
    }

    output = tmp_path / "syllabus.docx"
    build_syllabus_docx(syllabus, output)

    document = Document(output)
    assert document.tables[0].cell(0, 1).text == "2026-2027"
    assert document.tables[0].cell(1, 1).text == "Climate Change Law and Policy"
    assert document.tables[1].cell(0, 1).text == "M. Modele Syllabus"
    assert document.tables[5].cell(1, 1).text == "Evaluate climate policy."
    assert document.tables[6].cell(1, 0).text == "CLO 1: Analyse climate law."
    assert document.tables[7].cell(1, 0).text == "1"
    assert (
        document.tables[7].cell(1, 3).text
        == "Pre-class learning activities:\nRead chapter 1\n\nAssessments:\nShort quiz"
    )
    assert "A. Author" in document.tables[8].cell(0, 1).text
    assert document.tables[9].cell(1, 1).text == "Policy brief"
    assert "☒" in document.tables[10].cell(1, 1).text
    assert "Oral follow-up discussion" in document.tables[10].cell(1, 1).text
    assert document.tables[11].cell(0, 0).text == "Assessment type: Policy brief"
    assert document.tables[13].cell(0, 1).text == "23/07/2026"


def test_expands_outcome_rows_and_accepts_legacy_string_lists(tmp_path) -> None:
    syllabus = {
        "courseTitle": "Environmental Policy",
        "courseCode": "SCEN-220",
        "academicYear": "2026-2027",
        "content": {
            "learningOutcomes": {
                "plos": [f"PLO {index}: Outcome {index}" for index in range(1, 8)],
            }
        },
    }
    output = tmp_path / "expanded-syllabus.docx"

    build_syllabus_docx(syllabus, output)

    document = Document(output)
    assert len(document.tables[5].rows) == EXPANDED_PLO_TABLE_ROW_COUNT
    assert document.tables[5].cell(6, 1).text == "Outcome 6"
    assert document.tables[5].cell(7, 1).text == "Outcome 7"


def test_keeps_legacy_administrative_contact_text_readable(tmp_path) -> None:
    syllabus = {
        "courseTitle": "Environmental Policy",
        "courseCode": "SCEN-220",
        "academicYear": "2026-2027",
        "content": {
            "contacts": {
                "administrativeContact": LEGACY_CONTACT
            }
        },
    }
    output = tmp_path / "legacy-contact-syllabus.docx"

    build_syllabus_docx(syllabus, output)

    document = Document(output)
    assert document.tables[2].cell(0, 1).text.strip() == LEGACY_CONTACT
